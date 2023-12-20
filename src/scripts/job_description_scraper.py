import uuid

import requests
from bs4 import BeautifulSoup
# from docx.document import Document

from docx import Document

from src.utils import get_file_name

jd_path = "/Users/tobialao/Desktop/Software_Projects/msc_project/hiring_llm_dss_engine/resources/dataset/job_description"



def scrape_job_description_data(page_link: str, index=None):
    doc = Document()
    result = requests.get(page_link).text
    soup = BeautifulSoup(result, 'html.parser')
    job_title_tag = soup.find("h1", attrs={'class': 'masthead-title'})
    parent = soup.find("div", attrs={'class': 'content has-standfirst'})
    title = get_file_name(page_link)
    if job_title_tag is not None and job_title_tag.get_text():
        doc.core_properties.title = job_title_tag.get_text()
        doc.add_heading(f'Job Title: {job_title_tag.get_text()}', level=1)
        title = job_title_tag.get_text().strip().replace(" ", "_")
    if parent is not None:
        children = parent.findChildren(["p", "ul", "h2"])
        for tag in children:
            if tag.get('class') is None or 'anchor' in tag.get('class'):
                if tag.name == "p":
                    doc.add_paragraph(tag.get_text())
            if tag.name == "h2":
                # Break operation when it gets to irrelevant sections
                if tag.get_text().lower() in ['employers', 'advertisement', 'work experience',
                                              'related case studies',
                                              'professional development']:
                    break
                    # continue
                # elif tag.get_text().lower() in ['skills']:
                #     break
                doc.add_heading(tag.get_text(), 2)
            elif tag.name == "ul":
                li_responsibilities = tag.findChildren(["li"])
                for el in li_responsibilities:
                    doc.add_paragraph(
                        el.get_text(), style='List Bullet'
                    )

    doc.save(f"{jd_path}/{title}.docx")


if __name__ == "__main__":

    base_link = "https://www.prospects.ac.uk/job-profiles"

    page_paths = [
        # IT Roles
        "software-engineer",
        "application-analyst",
        "software-tester",
        "data-scientist",
        "network-engineer",

        # Finance Roles
        "chartered-accountant",
        "corporate-investment-banker",
        "pensions-manager",
        "tax-adviser",
        "pensions-manager",
        "risk-manager",

        # Engineering Roles
        "aerospace-engineer",
        "mechanical-engineer",
        "consulting-civil-engineer",
        "chemical-engineer",
        "structural-engineer",

        # Health Roles
        "cardiologist",
        "chiropractor",
        "dentist",
        "dietitian",
        "surgeon",
        "general-practice-doctor"
    ]

    for index, path in enumerate(page_paths):
        scrape_job_description_data(f"{base_link}/{path}", index)
